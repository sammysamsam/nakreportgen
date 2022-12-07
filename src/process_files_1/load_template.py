
import docx
import docx2pdf
import os

TEMPLATE = '..//TEMPLATE1.docx'
annots = []
paras = []
results = []

def load_template(t=None):
    global template
    global annots
    global paras
    global TEMPLATE

    if t is not None:
        TEMPLATE = t
    template = docx.Document(TEMPLATE)
    annots = []
    paras = []
    for para in template.paragraphs:
        paras.append(para)
    for table in template.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    annots.append(paragraph)

    # for r in annots:
    #     print(r.text)
def update_temp(key,value):
    if value is None:
        return

    for r in annots:
        if "#" + key + "#" in r.text:
            r.text = value
            r.text = r.text.replace("#" + key+"#",value)

    for r in paras:
        if "#" + key + "#" in r.text:
            r.text = r.text.replace("#" + key+"#",value)

def write_out_filled_template(file_name):
    template.save(file_name)


if __name__ == '__main__':
    load_template()




