
import docx
import docx2pdf
import pdf2docx
import copy

TEMPLATE = 'TEMPLATE2.docx'
annots = []
paras = []
results = []
tables = {}


def load_template(t=None):
    global template
    global annots
    global paras
    global tables

    global TEMPLATE
    if t is not None:
        TEMPLATE = t
    template = docx.Document(TEMPLATE)

    #find replacable values
    annots, paras = [], []
    for para in template.paragraphs:
        paras.append(para)
    for table in template.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    annots.append(paragraph)

    # find desired tables
    tables = {}
    for table in template.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text == "#table1":
                    tables["table1"] = table
                if cell.text == "#table2":
                    tables["table2"] = table
                if cell.text == "#table3":
                    tables["table3"] = table
                if cell.text == "#table4":
                    tables["table4"] = table
                if cell.text == "#table5":
                    tables["table5"] = table

    # template.save("tmp2.docx")

def update_temp(key,value):

    for r in annots:
        if "#" + key in r.text:
            if key in ["linearft","squareft"]:
                r.text = r.text.replace("#" + key, value)
            else:
                r.text = value
    for r in paras:
        if "#" + key in r.text:
            r.text = r.text.replace("#" + key, value)


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def remove_yes(t):
    keepgoing=True
    while keepgoing:
        keepgoing=False
        for r in t.rows:
            if r.cells[4].text == "No":
                keepgoing=True
                break
        if keepgoing:
            remove_row(t,r)
        else:
            break
    return t

def update_table(tablename, table_data):
    table = tables[tablename]
    remove_row(table, table.rows[0])
    if len(table_data) == 0:
        return
    for i in range(0,len(table_data)):
        table.add_row()  # ADD ROW HERE
        row = table_data[i]
        for ii in range(0,len(row)):
            if row[ii] is None:
                table.cell(i, ii).text = "NONE"
                continue
            table.cell(i, ii).text = row[ii]


def update_table_yes(tablename, table_data):
    table = tables[tablename]
    remove_row(table, table.rows[0])
    if len(table_data) == 0:
        return

    data = []
    for row in table_data:
        if row[4] == "Yes":
            data.append(row)
        elif "Asbestos" in row[4] and "Present" in row[4]:
            data.append(row)
    table_data = data

    for i in range(0,len(table_data)):
        table.add_row()  # ADD ROW HERE
        row = table_data[i]
        for ii in range(0,len(row)):
            if row[ii] is None:
                table.cell(i, ii).text = "NONE"
                continue
            table.cell(i, ii).text = row[ii]


def print_table(t):
    print("-_----------")
    for r in t.rows:
        x = []
        for c in r.cells:
            x.append(c.text.replace("\n",""))
        print(x)


def delete_columns(table, columns):
    # sort columns descending
    columns.sort(reverse=True)

    grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
    for ci in columns:
        for cell in table.column_cells(ci):
            cell._tc.getparent().remove(cell._tc)

        # Delete column reference.
        col_elem = grid[ci]
        grid.remove(col_elem)


def update_table_yes_three_columns(tablename, table_data):
    table = tables[tablename]
    if len(table_data) == 0:
        return

    update_table_yes(tablename, table_data)

    i = 0
    columns_remove = []
    for c in table.columns:
        t =  c.cells[0].text.replace(" ","").replace("\n","")

        if t in ["NumberofSamplesAnalyzed", 'AsbestosPresent','AssumedACM', 'SquareFeet', 'LinearFeet', 'Abated', 'ScopeofWorkApplied']:
            columns_remove.append(i)
        i += 1

    delete_columns(table, columns_remove)

def update_table_three_columns(tablename, table_data):
    table = tables[tablename]
    if len(table_data) == 0:
        return

    update_table(tablename, table_data)

    i = 0
    columns_remove = []
    for c in table.columns:
        t =  c.cells[0].text.replace(" ","").replace("\n","")

        if t in ["NumberofSamplesAnalyzed", 'AsbestosPresent','AssumedACM', 'SquareFeet', 'LinearFeet', 'Abated', 'ScopeofWorkApplied']:
            columns_remove.append(i)
        i += 1

    delete_columns(table, columns_remove)


def write_out_filled_template(file_name):
    template.save(file_name)


if __name__ == '__main__':
    load_template()
    update_temp("linearft", "hello")
    update_temp("squareft", "hello2")
    template.save("tmp2a.docx")